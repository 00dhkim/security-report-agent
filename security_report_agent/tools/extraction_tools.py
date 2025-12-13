import csv
import io
import subprocess
import os
import tempfile
from docx import Document
from typing import List, Dict, Any, Optional
from google.adk.tools import ToolContext

def _convert_doc_to_docx(doc_path):
    if doc_path.lower().endswith(".docx"):
        return doc_path
    output_dir = os.path.dirname(doc_path)
    subprocess.run(["libreoffice", "--headless", "--convert-to", "docx", doc_path, "--outdir", output_dir])
    return doc_path.replace(".doc", ".docx")

def _parse_docx_tables(path):
    doc = Document(path)
    all_text = []

    # 본문 텍스트
    for para in doc.paragraphs:
        all_text.append(para.text)

    # 테이블 텍스트
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip().replace(",", "") for cell in row.cells]
            all_text.append(",".join(row_text))
        all_text.append("")  # 테이블 간 구분을 위해 빈 줄 추가

    return "\n".join(all_text)

async def doc_report_parser(tool_context: ToolContext, report_path: Optional[str] = None) -> Dict[str, str]:
    """
    월간 보안 관제 doc/docx 보고서 파일을 파싱하여 일반 본문과 No,Src IP,Dest IP,Port,Count 컬럼 등을 포함한 문자열을 반환합니다.
    report_path가 제공되지 않으면 사용자가 업로드한 아티팩트(파일)를 사용합니다.
    """
    temp_file_path = None

    if not report_path:
        print("report_path not provided. Checking for uploaded artifacts...")
        try:
            artifacts = await tool_context.list_artifacts()
            if artifacts:
                # 가장 최근에 업로드된 아티팩트 사용 (리스트의 마지막 요소라고 가정)
                filename = artifacts[-1]
                print(f"Found artifact: {filename}")
                part = await tool_context.load_artifact(filename)
                
                if part and part.inline_data:
                    # 확장자 결정 (기본값 .doc)
                    _, ext = os.path.splitext(filename)
                    if not ext:
                        ext = ".doc"
                    
                    # 임시 파일 생성 및 데이터 쓰기
                    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
                        tmp.write(part.inline_data.data)
                        temp_file_path = tmp.name
                        report_path = tmp.name
                    print(f"Artifact saved to temporary file: {report_path}")
                else:
                    return {"error": "Artifact found but no data available."}
            else:
                return {"error": "No report_path provided and no artifacts found."}
        except Exception as e:
            return {"error": f"Failed to load artifact: {str(e)}"}

    if not report_path or not os.path.exists(report_path):
        return {"error": f"File not found: {report_path}"}
    
    try:
        docx_path = _convert_doc_to_docx(report_path)
        parsed_text = _parse_docx_tables(docx_path)
        
        # docx 변환 파일 삭제 (원본이 docx가 아니었던 경우에만 삭제)
        if docx_path != report_path and os.path.exists(docx_path):
            os.remove(docx_path)
            
        print(f"doc extraction result:\n{parsed_text[:50]}...")
        return {"csv_text": parsed_text}
        
    finally:
        # 아티팩트로부터 생성된 임시 파일 삭제
        if temp_file_path and os.path.exists(temp_file_path):
            os.remove(temp_file_path)

def csv_to_records(csv_text: str) -> Dict[str, List[Dict[str, Any]]]:
    """
    src_ip, dst_ip, dst_port, count 컬럼을 가진 CSV를 파싱하여 ReportRecord 리스트로 변환합니다.
    """
    print(f" csv_to_records(csv_text=...)")
    records = []
    # 문자열을 파일처럼 다루기 위해 io.StringIO 사용
    csv_file = io.StringIO(csv_text)
    reader = csv.DictReader(csv_file)
    for row in reader:
        records.append({
            "src_ip": row.get("src_ip"),
            "dst_ip": row.get("dst_ip"),
            "dst_port": int(row["dst_port"]) if row.get("dst_port") else None,
            "count": int(row["count"]) if row.get("count") else None,
        })
    return {"records": records}
